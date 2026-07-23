#!/usr/bin/env python3
"""Small, original DOCX helper for structural inspection and simple edits."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def ensure_distinct(source: Path, output: Path) -> None:
    if source.resolve() == output.resolve():
        raise SystemExit("Refusing to overwrite the input file; choose a different output path.")


def inspect_docx(path: Path) -> None:
    doc = Document(path)
    headings = []
    nonempty = 0
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            nonempty += 1
        if text and paragraph.style and paragraph.style.name.lower().startswith("heading"):
            headings.append({"paragraph": index, "style": paragraph.style.name, "text": text[:160]})

    tables = []
    for index, table in enumerate(doc.tables, start=1):
        tables.append(
            {
                "table": index,
                "rows": len(table.rows),
                "columns": max((len(row.cells) for row in table.rows), default=0),
            }
        )

    relationships = doc.part.rels.values()
    images = sum(1 for rel in relationships if "image" in rel.reltype)
    report = {
        "file": str(path.resolve()),
        "paragraphs": len(doc.paragraphs),
        "nonempty_paragraphs": nonempty,
        "tables": tables,
        "sections": len(doc.sections),
        "images": images,
        "headings": headings,
    }
    print(json.dumps(report, ensure_ascii=False, indent=2))


def add_content(doc: Document, text: str) -> None:
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif len(line) > 3 and line[:3].rstrip(".").isdigit() and line[1:3] in {". ", "、"}:
            doc.add_paragraph(line[3:].strip(), style="List Number")
        else:
            doc.add_paragraph(line)


def create_docx(input_path: Path, output: Path, title: str | None) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35

    if title:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(title)
        run.bold = True
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(22)
        paragraph.paragraph_format.space_after = Pt(18)

    add_content(doc, input_path.read_text(encoding="utf-8"))
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    inspect_docx(output)


def replace_text(source: Path, output: Path, old: str, new: str) -> None:
    ensure_distinct(source, output)
    doc = Document(source)
    replacements = 0
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            original = paragraph.text
            for run in paragraph.runs:
                run.text = ""
            paragraph.add_run(original.replace(old, new))
            replacements += original.count(old)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old in paragraph.text:
                        original = paragraph.text
                        for run in paragraph.runs:
                            run.text = ""
                        paragraph.add_run(original.replace(old, new))
                        replacements += original.count(old)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(json.dumps({"output": str(output.resolve()), "replacements": replacements}, ensure_ascii=False))


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser()
    sub = parser.add_subparsers(dest="command", required=True)

    inspect_cmd = sub.add_parser("inspect")
    inspect_cmd.add_argument("input", type=Path)

    create_cmd = sub.add_parser("create")
    create_cmd.add_argument("--input", required=True, type=Path)
    create_cmd.add_argument("--output", required=True, type=Path)
    create_cmd.add_argument("--title")

    replace_cmd = sub.add_parser("replace")
    replace_cmd.add_argument("input", type=Path)
    replace_cmd.add_argument("--output", required=True, type=Path)
    replace_cmd.add_argument("--old", required=True)
    replace_cmd.add_argument("--new", required=True)
    return parser


def main() -> None:
    args = build_parser().parse_args()
    if args.command == "inspect":
        inspect_docx(args.input)
    elif args.command == "create":
        create_docx(args.input, args.output, args.title)
    elif args.command == "replace":
        replace_text(args.input, args.output, args.old, args.new)


if __name__ == "__main__":
    main()
